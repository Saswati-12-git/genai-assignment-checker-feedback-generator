import os
import json
import streamlit as st
from groq import Groq
from PyPDF2 import PdfReader
from docx import Document
from difflib import SequenceMatcher
import re
import base64

# ==============================
# 1. Load API Key from Streamlit Secrets
# ==============================
api_key = st.secrets["groq"]["api_key"]

if not api_key:
    st.error("⚠️ GROQ API key not found in Streamlit secrets!")
    st.stop()

client = Groq(api_key=api_key)


# ==============================
# 2. Extract text helpers
# ==============================
def extract_text_from_file(file_path):
    text = ""
    lower = file_path.lower()

    if lower.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

    elif lower.endswith(".pdf"):
        try:
            reader = PdfReader(file_path)
            for page in reader.pages:
                text += (page.extract_text() or "") + "\n"
        except:
            return ""

    elif lower.endswith(".docx"):
        try:
            doc = Document(file_path)
            text = "\n".join(p.text for p in doc.paragraphs)
        except:
            return ""

    return text.strip()

def extract_text_from_uploaded(file):
    fname = file.name.lower()

    if fname.endswith(".txt"):
        return file.read().decode("utf-8")

    elif fname.endswith(".pdf"):
        reader = PdfReader(file)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages).strip()

    elif fname.endswith(".docx"):
        doc = Document(file)
        return "\n".join(p.text for p in doc.paragraphs).strip()

    return ""

# ==============================
# 3. Clean text for comparison
# ==============================
def clean_text(text):
    text = text.lower()
    text = re.sub(r'[^a-z0-9\s]', ' ', text)
    return re.sub(r'\s+', ' ', text).strip()

# ==============================
# 4. Plagiarism Checker
# ==============================
def check_plagiarism(text, folder="sample_essays"):
    text = clean_text(text)
    max_match = 0
    best_file = "No match detected"

    if not os.path.exists(folder):
        st.warning(f"⚠️ Folder '{folder}' not found. Create it and add reference files.")
        return 0, best_file

    for file in os.listdir(folder):
        path = os.path.join(folder, file)
        sample = extract_text_from_file(path)
        if not sample:
            continue

        sample = clean_text(sample)
        similarity = SequenceMatcher(None, text, sample).ratio()

        if similarity > max_match:
            max_match = similarity
            best_file = file

    return round(max_match * 100, 2), best_file

# ==============================
# 5. Streamlit UI
# ==============================
st.set_page_config(page_title="GenAI Assignment Checker", layout="wide")
st.title("🧠 GenAI Assignment Checker + Plagiarism Detector")

essay_text = st.text_area("✏️ Paste or write assignment here:", height=230)
uploaded = st.file_uploader("Or upload: TXT / PDF / DOCX", type=["txt", "pdf", "docx"])

if uploaded and not essay_text.strip():
    essay_text = extract_text_from_uploaded(uploaded)

if essay_text:
    st.success("✅ Text Loaded Successfully")

# ==============================
# 6. AI Evaluation Prompt
# ==============================
def build_prompt(text):
    return f"""
You are an academic writing evaluator. Analyze the text and respond ONLY in JSON:

{{
  "grammar_score": <1-10>,
  "coherence_score": <1-10>,
  "structure_score": <1-10>,
  "creativity_score": <1-10>,
  "overall_score": <1-100>,
  "summary": "<3-4 line summary of the writing>",
  "suggested_improvements": "<specific bullet point improvements>",
  "feedback": "<short professional feedback>"
}}

Text to evaluate:
{text}
"""

def evaluate_text(text):
    try:
        res = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": build_prompt(text)}],
            temperature=0.4
        )
        return json.loads(res.choices[0].message.content)
    except Exception as e:
        st.error(f"API Error: {e}")
        return None

# ==============================
# 7. Run Button
# ==============================
if st.button("🚀 Run Evaluation"):
    if not essay_text.strip():
        st.error("❗ Please enter or upload text first.")
    else:
        result = evaluate_text(essay_text)
        plag_percent, source_file = check_plagiarism(essay_text)

        if result:
            st.subheader("📊 Evaluation Results")
            st.json(result)

            # ✅ Download JSON Button
            json_data = json.dumps(result, indent=4)
            st.download_button("⬇️ Download Feedback (JSON)", json_data, file_name="feedback_report.json")

        st.subheader("🧾 Plagiarism Report")
        st.metric("Plagiarism Score", f"{plag_percent}%")
        st.write(f"Matched Source: **{source_file}**")

        if plag_percent > 30:
            st.error("⚠️ High plagiarism detected!")
        else:
            st.success("✅ Mostly original content!")

